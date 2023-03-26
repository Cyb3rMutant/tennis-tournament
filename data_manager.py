from abc import *
from os import listdir
from os.path import isdir, isfile
import pathlib
from docx import Document
import pandas as pd
import numpy as np
import re
import zipfile

DATA_PATH = f"{pathlib.Path(__file__).parent.resolve()}/data"

def check_filename(filename, check_name, check_extension):
    # get the filename and extension (remove path if exists)
    filename = filename.split("/")[-1]
    filename = filename.split(".")

    return (filename[0].lower() != check_name.lower() or filename[1].lower() != check_extension.lower())

class ModeNotFoundError(Exception):
    pass

class MethodNotFoundError(Exception):
    pass

class WrongFileExtensionError(Exception):
    pass

class BadFileError(Exception):
    pass

class UnformattedDocx(Exception):
    pass


class DataManager:
    '''
    The Data Manager class is used to find files in a data directory.

    Attributes:
        _mode: A string referring to the extension of the file(s) being read from, by default its "csv".
        _modes: An array of all the valid extensions for mode.
        _data_path: A string containing the path where the data files are located.
    '''

    def __init__(self, path=(pathlib.Path(__file__).parent.resolve())) -> None:
        self._modes = ["csv", "docx"]
        self._mode = "csv"
        self._data_path = path

    def set_mode(self, mode):
        '''
        Sets the Data Manager's mode to the input (if its valid).
        '''
        mode = mode.lower()
        if(mode in self._modes):
            self._mode = mode
        else:
            raise ModeNotFoundError(f'"{mode}" is not a valid mode, mode has to be one of the elements in {self._modes}.')

    def get_mode(self):
        '''
        Returns the current mode.
        '''
        return self._mode

    def set_data_path(self, path):
        '''
        Sets the data path to the input if its a real path.
        '''
        if(isdir(path)):
            if(path[-1] != "/"):
                path += "/"
            self._data_path = path
        else:
            raise FileNotFoundError(f'Path "{path}" not found')
        
    def get_data_path(self):
        '''
        Returns the data path.
        '''
        return self._data_path

    def get_data_files(self):
        '''
        Returns the names of all the files with the mode extension.
        '''
        files = [f for f in listdir(self._data_path) if (f.endswith(f".{self._mode}") and isfile(self._data_path+f))]
        return files

class DataExtractor(metaclass=ABCMeta):
    '''
    The Data Extractor class is used to extract data from csv and docx files (WIP)

    Attributes:
        _method: A string describing how the data files are inputted.
        _methods: An array of all the valid methods.
        _dm: A Data Manager Object used to retrieve all the files in a data directory.

    '''

    def __init__(self) -> None:
        self.dm = DataManager()
        self.dm.set_data_path(DATA_PATH)
        self._methods = ["upload", "path"]
        self._method = "upload"

    def set_method(self, method):
        '''
        Sets the method to the input (if its valid).
        '''
        method = method.lower()
        if(method in self._methods):
            self._method = method
        else:
            raise MethodNotFoundError(f'"{method}" is not a valid method, the method has to be one of the elements in {self._methods}.')

    def get_method(self):
        '''
        Returns the current method.
        '''
        return self._method

class DataExtractorCSV(DataExtractor):

    def __init__(self) -> None:
        super().__init__()
        self.dm.set_mode("csv")

    def get_players(self, files=[]):

        file_names = []
        file = None

        if(self._method.lower() == "upload"):
            if(files == []):
                raise FileNotFoundError("No player files passed to method")
            else:
                file_names = [file.filename for file in files]
                
        if(self._method.lower() == "path"):
            file_names = self.dm.get_data_files()


        players = {}
        for i in range(len(file_names)):
            doc = file_names[i]
            if(re.findall("PLAYERS*", doc)):

                if(self._method.lower() == "upload"):
                    file = files[i]
                elif(self._method.lower() == "path"):
                    file = f"{DATA_PATH}/{doc}"

                if(doc.endswith(".csv")):
                    df = pd.read_csv(file)
                else:
                    raise WrongFileExtensionError
                
                player_arr = df.to_numpy().T[0]
                key = doc.split(" ")[0].lower()
                players[key] = np.concatenate([df.columns.values, player_arr])

        return players
    
    def get_tournament_prizes(self, files=[]):

        file = None
        if(self._method.lower() == "upload"):
            if(files == []):
                raise FileNotFoundError("No files were inputted")
            else:
                file = files[0]  
                # check if the file name is prize money
                if(check_filename(file.filename, "PRIZE MONEY", "csv")):
                    return {} 
            
        if(self._method.lower() == "path"):
            file = f"{DATA_PATH}/PRIZE MONEY.csv"

        prizes = {}
        df = pd.read_csv(file)
        tournament = ""
        for row in df.to_numpy():
            if(not pd.isnull(row[0])):
                tournament = row[0].strip()
                prizes[tournament] = {}
            prizes[tournament][row[1]] = row[2]
        return prizes
    
    def get_ranking_points(self, files=[]):

        file = None
        if(self._method.lower() == "upload"):
            if(files == []):
                raise FileNotFoundError("No files were inputted")
            else:
                file = files[0]   
                # check if the file name is ranking points
                if(check_filename(file.filename, "RANKING POINTS", "csv")):
                    return {}

        if(self._method.lower() == "path"):
            file = f"{DATA_PATH}/RANKING POINTS.csv"

        points = {}
        df = pd.read_csv(file)
        for c, row in enumerate(df.to_numpy()):
            # not sure, theres an input mistake in csv data
            points[c+1] = row[0]
        return points
    
    def get_tournament_matches(self, tournament, files=[]):

        df = None
        file = None
        
        if(self._method.lower() == "upload"):
            if(files == []):
                raise FileNotFoundError("No round files passed to method")
            else:
                file_names = [file.filename for file in files]
                
        if(self._method.lower() == "path"):
            file_names = self.dm.get_data_files()

        tournament = tournament.upper()
        matches = {}
        for i in range(len(file_names)):
            doc = file_names[i]
            if(re.findall(f"{tournament} ROUND [0-9]+ *", doc.upper())):
                if(self._method.lower() == "upload"):
                    file = files[i]
                    # print(file)
                elif(self._method.lower() == "path"):
                    file = f"{DATA_PATH}/{doc}"
                df = pd.read_csv(file)
                tournament_round = doc.split(" ")[2]
                gender = doc.split(" ")[3].split(".")[0].lower()
                if(not matches.get(gender)):
                    matches[gender] = {}
                matches[gender][tournament_round] = df.iloc[:,:4].to_numpy()
                
        return matches
class DataExtractorDOCX(DataExtractor):

    def __init__(self) -> None:
        super().__init__()
        self.dm.set_mode("docx")
    
    def get_tournament_difficulty(self, files=[]):
        '''
        Returns a dictionary containing the tournament name and the tournament difficulty as key value pairs respectively.
        '''
        tournament_dict = {}
        file = None
        document = None

        if(self._method.lower() == "upload"):
            if(files == []):
                raise FileNotFoundError("No file was inputted")
            else:
                file = files[0]

                #VALIDATION VALIDATION VALIDATION VALIDATION VALIDATION VALIDATION VALIDATION VALIDATION
                if file.filename.lower().endswith('docx'):
                    try:
                        document = Document(file)
                    except:
                        raise BadFileError("File is a fake docx")
                else:
                    raise WrongFileExtensionError("File is not a docx")
                
        if(self._method.lower() == "path"):
            file = f"{DATA_PATH}/DEGREE OF DIFFICULTY PER TOURNAMENT.docx"
            document = Document(file)

        # go line by line through the document
        for i in range(len(document.paragraphs)):
            # for every line split the text to get the tournament name and difficulty
            tournament = document.paragraphs[i].text
            tournament = tournament.split(" – degree of difficulty ")
            # add the tournament name and difficulty to the tournament dictionary if they exist on that line
            if(len(tournament) > 1):
                tournament_dict[tournament[0]] = tournament[1]
        
        
        if tournament_dict == {}:
            raise UnformattedDocx("File is docx but contents are invalid")
        else:
            return tournament_dict
        

if(__name__ == "__main__"):

    def openFile(files, filename):
        fp = open(f"data/{filename}", 'rb')
        file = FileStorage(fp)
        files.append(file)

    from werkzeug.datastructures import FileStorage

    # dm = DataManager()
    # dm.set_mode("docx")
    # print(dm.get_mode())
    # dm.set_data_path(DATA_PATH)
    # print(dm.get_data_files())

    # document = Document(f"{DATA_PATH}/DEGREE OF DIFFICULTY PER TOURNAMENT.docx")
    # for p in document.paragraphs:
    #     print(p.text)

    de = DataExtractorCSV()
    # de.set_method("path")
    # print(de.get_tournament_matches("tac1"))
    # import pprint 
    # pp = pprint.PrettyPrinter(indent=4)
    # pp.pprint(de.get_tournament_matches("tac1"))
    # print(de.get_tournament_matches("tac1").get("men").get("5")[0][2])
    
    files = []
    openFile(files, "RANKING POINTS.csv")
    # openFile(files, "MALE PLAYERS.csv")
    # openFile(files, "FEMALE PLAYERS.csv")
    # openFile(files, "TAC1 ROUND 5 MEN.csv")
    # openFile(files, "TAC1 ROUND 2 LADIES.csv")
    # openFile(files, "TAC1 ROUND 3 LADIES.csv")
    # openFile(files, "PRIZE MONEY.csv")
    
    # print(de.get_tournament_prizes(files=files))
    print(de.get_ranking_points(files))


    
    